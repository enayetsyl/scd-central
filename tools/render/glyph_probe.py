#!/usr/bin/env python3
"""glyph_probe.py — WIP (tools/_wip). Establishes a render path's proven glyph set.

CD-018: "each render path proves its own glyph set empirically and records it in that path's
SMOKE.md". This is the instrument that produces that evidence for the docx path.

Three independent signals, because no single one is trustworthy:

  1. COVERAGE (static, fontTools) — does the font's cmap actually contain the codepoint?
     Definitive. A missing cmap entry means tofu, full stop.
  2. SUBSTITUTION (dynamic, pdffonts) — which fonts did LibreOffice actually embed? A font
     appearing that we did not ask for means LibreOffice fell back, i.e. our font lacked
     the glyph and something else covered it.
  3. RASTER (pdftoppm PNG) — for human eyeball verification. CD-014 requires eyeball
     confirmation for Arabic (RTL + joining), which no automated check can replace.

Deliberately NOT used as evidence: reading the codepoint back out of the .docx XML, or
pdftotext round-trip. A correct codepoint in the text layer can still render as a box — that
is precisely the failure this probe exists to catch (tools/_wip/STATE.md, CD-018).

Usage:
    python3 glyph_probe.py --fonts-dir <dir> [--out <dir>] [--set core|arabic|all]

Exit 0 = probe ran (findings are in the report, not the exit code).
Exit 1 = probe could not run (missing tool or font).
"""
import argparse
import os
import re
import shutil
import subprocess
import sys
import unicodedata
from pathlib import Path

# ---- the test set, per CD-018 + tools/_wip/STATE.md -------------------------
CORE = [
    ("Bengali numerals", "০১২৩৪৫৬৭৮৯"),
    ("Bengali basic", "আমার বাংলা বই"),
    ("যুক্তবর্ণ (conjuncts)", "হ্ম ক্ষ ঞ্জ ন্দ্র স্ট্র"),
    ("কারচিহ্ন (vowel signs)", "কা কি কী কু কূ কৃ কে কৈ কো কৌ"),
    ("special marks", "অংশ দুঃখ চাঁদ"),
    ("legend glyphs", "🔴 🟦 ★ ↑ ↓"),
    ("tier-3 WATCH", "— …"),
    ("Latin + digits", "Class 5 — 100 marks"),
]
ARABIC = [
    ("Arabic ayah (CD-014)", "بِسْمِ اللَّهِ الرَّحْمَٰنِ الرَّحِيمِ"),
    ("Arabic bare", "الحمد لله"),
]


def need(tool):
    if shutil.which(tool) is None:
        sys.exit(f"ERROR: required tool not found: {tool}")


def load_fonts(fonts_dir):
    files = sorted([p for p in Path(fonts_dir).glob("**/*")
                    if p.suffix.lower() in {".ttf", ".otf", ".ttc"}])
    if not files:
        sys.exit(f"ERROR: no font files (.ttf/.otf/.ttc) under {fonts_dir}")
    return files


def coverage(font_path, chars):
    """Return the set of characters the font's cmap does NOT cover."""
    from fontTools.ttLib import TTFont, TTCollection
    fonts = []
    if font_path.suffix.lower() == ".ttc":
        fonts = list(TTCollection(str(font_path)).fonts)
    else:
        fonts = [TTFont(str(font_path), fontNumber=0, lazy=True)]
    covered = set()
    for f in fonts:
        try:
            for table in f["cmap"].tables:
                covered |= set(table.cmap.keys())
        except Exception:
            pass
    missing = set()
    for ch in chars:
        if unicodedata.category(ch) in {"Zs", "Cc"}:
            continue
        if ord(ch) not in covered:
            missing.add(ch)
    return missing


def build_docx(rows, font_name, out_path):
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(14)
    # python-docx does not set the complex/east-asian font runs; do it explicitly or
    # LibreOffice picks its own default for non-Latin script runs and the test is void.
    rpr = style.element.get_or_add_rPr().get_or_add_rFonts()
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rpr.set(qn(attr), font_name)
    doc.add_paragraph(f"glyph_probe — font: {font_name}")
    for label, text in rows:
        doc.add_paragraph(f"{label}: {text}")
    doc.save(out_path)


def to_pdf(docx_path, outdir):
    subprocess.run(["soffice", "--headless", "--convert-to", "pdf",
                    "--outdir", str(outdir), str(docx_path)],
                   capture_output=True, text=True, timeout=180)
    pdf = Path(outdir) / (Path(docx_path).stem + ".pdf")
    return pdf if pdf.exists() else None


def embedded_fonts(pdf):
    r = subprocess.run(["pdffonts", str(pdf)], capture_output=True, text=True)
    names = []
    for line in r.stdout.splitlines()[2:]:
        if line.strip():
            names.append(re.sub(r"^[A-Z]{6}\+", "", line.split()[0]))
    return sorted(set(names))


def rasterise(pdf, outdir):
    subprocess.run(["pdftoppm", "-png", "-r", "150", str(pdf),
                    str(Path(outdir) / "page")], capture_output=True, text=True)
    return sorted(Path(outdir).glob("page*.png"))


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--fonts-dir", required=True)
    ap.add_argument("--out", default="probe_out")
    ap.add_argument("--set", choices=["core", "arabic", "all"], default="all")
    ap.add_argument("--corpus", nargs="*", default=[],
                    help="Real documents to probe. The corpus is the honest test: a synthetic "
                         "string only proves what someone thought to type.")
    args = ap.parse_args()

    need("soffice"); need("pdftoppm"); need("pdffonts")
    rows = {"core": CORE, "arabic": ARABIC, "all": CORE + ARABIC}[args.set]
    for path in args.corpus:
        text = Path(path).read_text(encoding="utf-8")
        # Every distinct non-ASCII, non-space character actually used in the document.
        chars = "".join(sorted({c for c in text
                                if ord(c) > 0x7F and unicodedata.category(c) not in {"Zs", "Cc"}}))
        rows = rows + [(f"CORPUS {Path(path).name}", chars)]
    outdir = Path(args.out); outdir.mkdir(parents=True, exist_ok=True)
    fonts = load_fonts(args.fonts_dir)

    print(f"glyph_probe — {len(fonts)} font file(s) under {args.fonts_dir}\n")

    print("== 1. COVERAGE (font cmap — definitive) ==")
    for fp in fonts:
        for label, text in rows:
            miss = coverage(fp, text)
            mark = "ok  " if not miss else "MISS"
            detail = "" if not miss else "  missing: " + " ".join(
                f"U+{ord(c):04X}({c})" for c in sorted(miss))
            print(f"  [{mark}] {fp.name:28} {label}{detail}")
    print()

    print("== 2. SUBSTITUTION + 3. RASTER (docx -> LibreOffice -> pdftoppm) ==")
    for fp in fonts:
        from fontTools.ttLib import TTFont
        try:
            name = TTFont(str(fp), fontNumber=0, lazy=True)["name"].getDebugName(1)
        except Exception:
            name = fp.stem
        d = outdir / fp.stem
        d.mkdir(parents=True, exist_ok=True)
        docx = d / f"probe_{fp.stem}.docx"
        build_docx(rows, name, docx)
        pdf = to_pdf(docx, d)
        if pdf is None:
            print(f"  [FAIL] {fp.name}: LibreOffice produced no PDF")
            continue
        emb = embedded_fonts(pdf)
        unexpected = [e for e in emb if name.replace(" ", "") not in e.replace(" ", "")]
        pngs = rasterise(pdf, d)
        print(f"  {fp.name}  (requested: {name})")
        print(f"    embedded: {', '.join(emb) or '(none)'}")
        if unexpected:
            print(f"    SUBSTITUTION: {', '.join(unexpected)} — LibreOffice fell back; "
                  "the requested font did not cover something")
        print(f"    raster: {', '.join(p.name for p in pngs) or '(none)'} in {d}")
    print("\nEYEBALL REQUIRED: open the PNGs. CD-014 needs RTL + joining confirmed by a human "
          "for Arabic; no automated signal above substitutes for that.")


if __name__ == "__main__":
    main()
