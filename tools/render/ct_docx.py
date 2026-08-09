#!/usr/bin/env python3
"""ct_docx.py — class-test Markdown to .docx, for the SCD print path.

Authored against the two reference class tests in `reference/` (C5 Bangla, ch.19 and ch.20).
Font family is configuration, not a hardcoded assumption (CD-018: each render path proves its
own glyph set — so the path must be able to change fonts and re-prove).

Glyph handling — the point of this tool:
  The body font (Noto Serif Bengali) covers all Bengali but no symbol glyphs. Rather than
  hardcode a substitution list, the generator ASKS EACH FONT WHAT IT COVERS (fontTools cmap)
  and routes every character to a font that actually has it:

      body font   -> Bengali, Latin, em-dash, ellipsis
      symbol font -> characters the body font lacks but the symbol font has (✓, →, ★, ↑, ↓)
      arabic font -> Arabic script (CD-014; RTL + joining proven in SMOKE.md)

  Characters no configured font covers are reported as UNRESOLVED and would tofu. The tool
  refuses to claim success while any remain — that is the failure this whole path exists to
  prevent (UP-001).

Template substitutions (Principal ruling 2026-08-09, finding R-1):
      ✅  ->  ✓            (rendered via the symbol run)
      ⚠   ->  "দ্রষ্টব্য:"   (Bengali, no symbol needed)
      ✓   kept, routed to the symbol font
  → is routed to the symbol font by the same mechanism (see SMOKE.md, noted as an extension
  of the ✓ ruling rather than a separate one).

Usage:
    python3 ct_docx.py <input.md> [-o out.docx] [--body-font ...] [--symbol-font ...]
                       [--arabic-font ...] [--fonts-dir ...] [--strict]
Exit 0 = written. Exit 1 = unresolved glyphs (with --strict) or input error.
"""
import argparse
import re
import sys
import unicodedata
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, Inches

HERE = Path(__file__).resolve().parent

DEFAULT_BODY = "Noto Serif Bengali"
DEFAULT_SYMBOL = "DejaVu Sans"
DEFAULT_ARABIC = "Noto Naskh Arabic"

# R-1 ruling: applied to source text before any layout decision.
SUBSTITUTIONS = [("\u2705", "\u2713"), ("\u26a0\ufe0f", "দ্রষ্টব্য:"), ("\u26a0", "দ্রষ্টব্য:")]

ARABIC_RANGES = [(0x0600, 0x06FF), (0x0750, 0x077F), (0xFB50, 0xFDFF), (0xFE70, 0xFEFF)]


def is_arabic(ch):
    c = ord(ch)
    return any(a <= c <= b for a, b in ARABIC_RANGES)


def cmap_of(path):
    from fontTools.ttLib import TTFont
    f = TTFont(str(path), fontNumber=0, lazy=True)
    cov = set()
    for t in f["cmap"].tables:
        cov |= set(t.cmap.keys())
    return cov


def build_coverage(fonts_dir, body, symbol, arabic):
    """Map font-name -> covered codepoints, for every font we can find on disk."""
    cov = {}
    d = Path(fonts_dir)
    for p in sorted(d.glob("*.ttf")) + sorted(d.glob("*.otf")):
        from fontTools.ttLib import TTFont
        try:
            name = TTFont(str(p), fontNumber=0, lazy=True)["name"].getDebugName(1)
            cov.setdefault(name, set()).update(cmap_of(p))
        except Exception:
            continue
    # The symbol font is a system font, not vendored; probe it via fontconfig.
    if symbol not in cov:
        import subprocess
        r = subprocess.run(["fc-match", "-f", "%{file}", symbol],
                           capture_output=True, text=True)
        if r.stdout.strip() and Path(r.stdout.strip()).exists():
            cov[symbol] = cmap_of(Path(r.stdout.strip()))
    return cov


class Router:
    """Decides which configured font each character is rendered in."""

    def __init__(self, cov, body, symbol, arabic):
        self.cov, self.body, self.symbol, self.arabic = cov, body, symbol, arabic
        self.unresolved = {}

    def font_for(self, ch):
        if ch in "\n\t " or unicodedata.category(ch) in {"Zs", "Cc"}:
            return self.body
        if is_arabic(ch) and ord(ch) in self.cov.get(self.arabic, set()):
            return self.arabic
        if ord(ch) in self.cov.get(self.body, set()):
            return self.body
        if ord(ch) in self.cov.get(self.symbol, set()):
            return self.symbol
        if ord(ch) in self.cov.get(self.arabic, set()):
            return self.arabic
        self.unresolved[ch] = self.unresolved.get(ch, 0) + 1
        return self.body

    def segments(self, text):
        """Split text into (font, run_text) preserving order."""
        out = []
        for ch in text:
            f = self.font_for(ch)
            if out and out[-1][0] == f:
                out[-1][1] += ch
            else:
                out.append([f, ch])
        return out


def set_run_font(run, name, size=None, bold=None, italic=None):
    run.font.name = name
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    # python-docx only sets w:ascii; without the other three, LibreOffice picks its own font
    # for non-Latin runs and every glyph guarantee above is void.
    rf = run._element.get_or_add_rPr().get_or_add_rFonts()
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rf.set(qn(attr), name)


INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*)")


def add_inline(par, text, router, size, base_bold=False):
    """Emit text with **bold** / *italic*, each split further by font coverage."""
    for piece in INLINE.split(text):
        if not piece:
            continue
        bold, italic = base_bold, False
        if piece.startswith("**") and piece.endswith("**"):
            piece, bold = piece[2:-2], True
        elif piece.startswith("*") and piece.endswith("*"):
            piece, italic = piece[1:-1], True
        for font, chunk in router.segments(piece):
            set_run_font(par.add_run(chunk), font, size, bold, italic)


def apply_substitutions(md_text):
    for a, b in SUBSTITUTIONS:
        md_text = md_text.replace(a, b)
    return md_text


def convert(md_text, router, out_path, base_size=12):
    """md_text must already have been through apply_substitutions()."""
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.6)
        s.left_margin = s.right_margin = Inches(0.7)

    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if not line.strip():
            i += 1
            continue

        if re.fullmatch(r"-{3,}", line.strip()):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            pr = p._p.get_or_add_pPr()
            from docx.oxml import OxmlElement
            pbdr = OxmlElement("w:pBdr"); bot = OxmlElement("w:bottom")
            bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "6")
            bot.set(qn("w:space"), "1"); bot.set(qn("w:color"), "999999")
            pbdr.append(bot); pr.append(pbdr)
            i += 1
            continue

        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level, text = len(m.group(1)), m.group(2)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level <= 2 else WD_ALIGN_PARAGRAPH.LEFT
            add_inline(p, text, router, {1: 15, 2: 13.5}.get(level, 12.5), base_bold=True)
            i += 1
            continue

        if line.lstrip().startswith("|"):
            block = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                block.append(lines[i].strip())
                i += 1
            rows = [[c.strip() for c in r.strip("|").split("|")] for r in block
                    if not re.fullmatch(r"\|[\s:\-|]+\|", r)]
            if rows:
                t = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
                t.style = "Table Grid"
                for ri, row in enumerate(rows):
                    for ci, cell in enumerate(row):
                        par = t.cell(ri, ci).paragraphs[0]
                        add_inline(par, cell, router, 10.5, base_bold=(ri == 0))
            continue

        if line.lstrip().startswith(">"):
            text = re.sub(r"^\s*>\s?", "", line)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.35)
            add_inline(p, text, router, base_size)
            i += 1
            continue

        m = re.match(r"^\s*[-*]\s+(.*)$", line)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, m.group(1), router, base_size)
            i += 1
            continue

        p = doc.add_paragraph()
        add_inline(p, line, router, base_size)
        i += 1

    doc.save(out_path)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("input")
    ap.add_argument("-o", "--output")
    ap.add_argument("--fonts-dir", default=str(HERE / "fonts"))
    ap.add_argument("--body-font", default=DEFAULT_BODY)
    ap.add_argument("--symbol-font", default=DEFAULT_SYMBOL)
    ap.add_argument("--arabic-font", default=DEFAULT_ARABIC)
    ap.add_argument("--strict", action="store_true",
                    help="exit 1 if any character has no covering font")
    args = ap.parse_args()

    src = Path(args.input)
    if not src.exists():
        sys.exit(f"ERROR: input not found: {src}")
    out = Path(args.output) if args.output else src.with_suffix(".docx")

    cov = build_coverage(args.fonts_dir, args.body_font, args.symbol_font, args.arabic_font)
    for label, name in (("body", args.body_font), ("symbol", args.symbol_font)):
        if name not in cov:
            sys.exit(f"ERROR: {label} font '{name}' not found in {args.fonts_dir} or on the system")

    raw = src.read_text(encoding="utf-8")
    text = apply_substitutions(raw)
    router = Router(cov, args.body_font, args.symbol_font, args.arabic_font)
    convert(text, router, out)

    print(f"ct_docx: {src.name} -> {out}")
    print(f"  fonts: body={args.body_font} · symbol={args.symbol_font} · arabic={args.arabic_font}")
    applied = [f"{a}->{b}" for a, b in SUBSTITUTIONS if a in raw]
    if applied:
        print("  substitutions applied (R-1): " + " · ".join(applied))
    # Report against the SUBSTITUTED text — the raw source still contains characters that
    # were deliberately replaced, and reporting those would be a false alarm.
    routed = sorted({c for c in set(text)
                     if router.font_for(c) == args.symbol_font})
    if routed:
        print("  routed to symbol font: " + " ".join(f"U+{ord(c):04X}({c})" for c in routed))
    if router.unresolved:
        print("  UNRESOLVED (would tofu): " + " ".join(
            f"U+{ord(c):04X}({c})x{n}" for c, n in sorted(router.unresolved.items())))
        if args.strict:
            sys.exit(1)
    else:
        print("  UNRESOLVED: none — every character has a covering font")


if __name__ == "__main__":
    main()
